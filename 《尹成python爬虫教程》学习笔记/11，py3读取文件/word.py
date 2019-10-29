# 读取docx中的文本代码示例
import docx

# 获取文档对象
file = docx.Document(r"123.docx")
for para in file.paragraphs:
    print(para.text)

# print("段落数:" + str(len(file.paragraphs)))  # 段落数，每个回车隔离一段

# 输出段落编号及段落内容
# for i in range(len(file.paragraphs)):
#     print("第" + str(i) + "段的内容是：" + file.paragraphs[i].text)
